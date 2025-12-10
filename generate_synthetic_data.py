import os
import random
from datetime import datetime
from faker import Faker
from PIL import Image, ImageDraw, ImageFont, ImageFilter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document

# --- Configuration ---
OUTPUT_DIR = "./data/synthetic_samples"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Faker for US names (Asylum Seekers)
fake = Faker("en_US")

# --- 1. HK Localization Helpers ---


def get_hk_name():
    """Generates a typical HK name: English First + HK Surname."""
    # Common HK surnames (Romanized Cantonese)
    hk_surnames = [
        "CHAN",
        "WONG",
        "LEE",
        "CHEUNG",
        "LAU",
        "LAM",
        "IP",
        "NG",
        "HO",
        "YEUNG",
        "CHENG",
        "TANG",
    ]
    surname = random.choice(hk_surnames)
    first_name = fake.first_name()  # English first name
    return f"{first_name} {surname}"


def get_hk_address():
    """Generates a realistic HK address structure."""
    districts = [
        "Sha Tin",
        "Mong Kok",
        "Wan Chai",
        "North Point",
        "Tuen Mun",
        "Kwun Tong",
        "Central",
        "Causeway Bay",
    ]
    streets = [
        "King's Road",
        "Nathan Road",
        "Hennessy Road",
        "Tai Po Road",
        "Des Voeux Road",
        "Queen's Road",
    ]
    estates = ["Garden City", "City One", "Tai Koo Shing", "Whampoa Garden"]

    structure = random.choice(["street", "estate"])

    flat = f"Flat {random.choice(['A','B','C','D'])}, {random.randint(1, 40)}/F"
    district = random.choice(districts)

    if structure == "street":
        street = random.choice(streets)
        return f"{flat}, {random.randint(1, 900)} {street}, {district}, Hong Kong"
    else:
        estate = random.choice(estates)
        block = f"Block {random.randint(1, 10)}"
        return f"{flat}, {block}, {estate}, {district}, Hong Kong"


# --- 2. Data Structures ---


class Person:
    def __init__(self, role="Primary"):
        self.role = role
        self.gender = random.choice(["M", "F"])
        # US Names for Asylum Seekers
        self.first_name = (
            fake.first_name_male() if self.gender == "M" else fake.first_name_female()
        )
        self.last_name = fake.last_name()

        # Format: SURNAME, Firstname (Standard for official docs)
        self.name = f"{self.last_name.upper()}, {self.first_name}"

        self.dob = fake.date_of_birth(
            minimum_age=18 if role == "Primary" else 2,
            maximum_age=60 if role == "Primary" else 17,
        )
        self.nationality = random.choice(["IDN", "PHL", "VNM", "NPL", "PAK"])
        self.recognizance_no = (
            f"HK-{random.randint(1000000, 9999999)}-{random.choice(['A','R'])}"
        )
        self.prev_recognizance_no = (
            f"HK-{random.randint(1000000, 9999999)}-{random.choice(['A','R'])}"
        )


class Family:
    def __init__(self, family_id):
        self.id = family_id
        self.primary = Person(role="Primary")

        # Dependents share the primary's last name
        self.dependents = []
        for _ in range(random.randint(1, 2)):
            dep = Person(role="Dependent")
            dep.last_name = self.primary.last_name
            dep.name = f"{dep.last_name.upper()}, {dep.first_name}"
            self.dependents.append(dep)

        # Shared Family Details
        self.address = get_hk_address()  # Use HK Address
        self.rent_amt = random.randint(4500, 12000)
        self.rent_date = random.randint(1, 28)
        self.landlord = get_hk_name()  # Use HK Name for Landlord
        self.case_worker = get_hk_name()  # HK Name for Caseworker too


# --- 3. Font Utilities ---
def get_pil_font(size, bold=False):
    """Safe font loader for JPG generation."""
    try:
        font_path = "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"
        return ImageFont.truetype(font_path, size)
    except IOError:
        try:
            font_path = "Arial Bold.ttf" if bold else "Arial.ttf"
            return ImageFont.truetype(font_path, size)
        except IOError:
            return ImageFont.load_default()


# --- 4. Document Generators ---


def generate_recognizance_jpg(person, family_address):
    """Generates a messy scanned JPG form for ONE person."""
    W, H = 1000, 1400
    img = Image.new("RGB", (W, H), (255, 255, 250))
    draw = ImageDraw.Draw(img)

    header_font = get_pil_font(36, bold=True)
    label_font = get_pil_font(24)
    data_font = get_pil_font(24, bold=True)
    small_font = get_pil_font(20)

    # Header
    draw.text((320, 50), "IMMIGRATION DEPARTMENT", font=header_font, fill="black")
    draw.text((380, 100), "RECOGNIZANCE (Form 8)", font=header_font, fill="black")

    start_y = 200
    line_height = 70
    valid_until = fake.future_date(end_date="+30d")
    report_day = random.choice(["Monday", "Wednesday", "Friday"])

    data = [
        ("Serial No:", f"S-{random.randint(10000, 99999)}"),
        ("Recognizance No:", person.recognizance_no),
        ("Previous Rec No:", person.prev_recognizance_no),
        ("Name of Holder:", person.name),
        ("Nationality:", person.nationality),
        ("Sex:", person.gender),
        ("Date of Birth:", person.dob.strftime("%d-%b-%Y")),
        ("Residential Address:", family_address[:35]),
        ("", family_address[35:]),
        ("Valid Until:", valid_until.strftime("%d-%b-%Y")),
        ("Reporting Condition:", f"Report to Shatin Imm. Office"),
        ("", f"Every 4 weeks on {report_day}"),
    ]

    for label, value in data:
        draw.text((50, start_y), label, font=label_font, fill="black")
        draw.text((350, start_y), value, font=data_font, fill="darkblue")
        start_y += line_height

    # Photo Box
    draw.rectangle([750, 220, 920, 420], outline="black", width=3)
    draw.text((790, 310), "PHOTO", font=small_font, fill="gray")

    # Stamp
    draw.ellipse([100, 1100, 350, 1300], outline="red", width=5)
    draw.text((160, 1180), "OFFICIAL\nUSE ONLY", font=header_font, fill="red")

    # Scan Artifacts
    img = img.rotate(
        random.uniform(-0.8, 0.8),
        resample=Image.BICUBIC,
        expand=True,
        fillcolor="white",
    )
    img = img.filter(ImageFilter.GaussianBlur(0.7))

    filename = f"recognizance_{person.recognizance_no}.jpg"
    img.save(os.path.join(OUTPUT_DIR, filename))
    print(f"Generated JPG: {filename}")


def generate_tenancy_pdf(family):
    """Generates a structured PDF contract for the whole family."""
    filename = f"tenancy_{family.id}.pdf"
    filepath = os.path.join(OUTPUT_DIR, filename)

    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height - 60, "RESIDENTIAL TENANCY AGREEMENT")

    c.setFont("Helvetica", 12)
    y = height - 120
    line_height = 25

    all_tenants = f"{family.primary.name}"
    if family.dependents:
        all_tenants += " and " + ", ".join([d.name for d in family.dependents])

    lease_start = datetime.now().replace(day=1)
    lease_end = lease_start.replace(year=lease_start.year + 2)

    content = [
        ("Date:", datetime.now().strftime("%Y-%m-%d")),
        ("Landlord:", family.landlord),
        ("Tenant(s):", all_tenants),
        ("Premises:", family.address),
        (" ", ""),
        ("1. RENT", ""),
        (
            f"   The Tenant agrees to pay the monthly rent of HKD ${family.rent_amt:,}.",
            "",
        ),
        (f"   Rent is due on the {family.rent_date}th day of each calendar month.", ""),
        (" ", ""),
        ("2. TERM", ""),
        (
            f"   The lease term is for 24 months commencing on {lease_start.strftime('%d-%b-%Y')}.",
            "",
        ),
        (f"   Expiry Date: {lease_end.strftime('%d-%b-%Y')}.", ""),
    ]

    for left, right in content:
        c.drawString(50, y, left)
        if right:
            c.drawString(200, y, right)
        y -= line_height

    y -= 50
    c.line(50, y, 250, y)
    c.drawString(50, y - 15, "Signature of Landlord")
    c.line(350, y, 550, y)
    c.drawString(350, y - 15, "Signature of Tenant")

    c.save()
    print(f"Generated PDF: {filename}")


def generate_casenote_docx(family):
    """Generates a Word Doc with narrative events."""
    filename = f"casenote_{family.id}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    doc.add_heading("Social Worker Case Note", 0)

    table = doc.add_table(rows=3, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "File Ref:"
    cells[1].text = f"SW-{family.id}-{random.randint(100,999)}"
    cells = table.rows[1].cells
    cells[0].text = "Primary Client:"
    cells[1].text = family.primary.name
    cells = table.rows[2].cells
    cells[0].text = "Caseworker:"
    cells[1].text = family.case_worker

    doc.add_heading(
        f"Event 1: Home Visit ({fake.date_this_month().strftime('%Y-%m-%d')})", level=1
    )
    p1 = doc.add_paragraph()
    p1.add_run("Summary: ").bold = True
    p1.add_run(f"Conducted a routine home visit at {family.address}. ")
    p1.add_run(
        f"The primary client, {family.primary.first_name}, was present along with {len(family.dependents)} dependent(s). "
    )
    p1.add_run("The living environment appeared safe but cluttered. ")

    doc.add_heading(
        f"Event 2: Office Interview ({fake.date_this_month().strftime('%Y-%m-%d')})",
        level=1,
    )
    p2 = doc.add_paragraph()
    p2.add_run("Summary: ").bold = True
    p2.add_run(f"Client visited the center regarding rent arrears. ")
    p2.add_run(f"Landlord {family.landlord} has issued a verbal warning. ")
    p2.add_run(f"Current rent is HKD ${family.rent_amt}. Client requested assistance.")

    doc.save(filepath)
    print(f"Generated DOCX: {filename}")


# --- Main Execution ---
if __name__ == "__main__":
    print(f"Generating cohesive case files in {OUTPUT_DIR}...")

    # 1. FIXED: Argument name changed from 'id' to 'family_id' to match __init__
    families = [Family(family_id=101), Family(family_id=102)]

    for fam in families:
        print(f"\n--- Processing Family {fam.id}: {fam.primary.name} ---")
        generate_recognizance_jpg(fam.primary, fam.address)
        for dep in fam.dependents:
            generate_recognizance_jpg(dep, fam.address)
        generate_tenancy_pdf(fam)
        generate_casenote_docx(fam)

    print("\nBatch generation complete.")
